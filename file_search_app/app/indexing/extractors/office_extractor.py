"""Extração de texto de documentos Office (DOCX via python-docx, XLSX via openpyxl)."""

from __future__ import annotations

import logging
from pathlib import Path

from app.indexing.extractors.base import BaseExtractor, ExtractionResult
from app.services.settings_service import AppConfig

logger = logging.getLogger(__name__)

try:
    import docx

    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    import openpyxl

    _XLSX_OK = True
except ImportError:
    _XLSX_OK = False

# Limites defensivos para planilhas gigantes.
_MAX_XLSX_ROWS_PER_SHEET = 5000
_MAX_XLSX_CELLS = 200_000


class DocxExtractor(BaseExtractor):
    """Extrai parágrafos e tabelas de arquivos DOCX."""

    extensions = ("docx",)

    def is_available(self) -> bool:
        return _DOCX_OK

    def extract(self, path: Path, config: AppConfig) -> ExtractionResult:
        if not _DOCX_OK:
            return ExtractionResult(
                status="metadata_only",
                error="python-docx não instalado (pip install python-docx)",
            )
        try:
            document = docx.Document(str(path))
            parts = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return ExtractionResult(text="\n".join(parts)).truncated()
        except Exception as exc:
            logger.warning("Falha ao ler DOCX %s: %s", path, exc)
            return ExtractionResult(status="error", error=f"DOCX ilegível: {exc}")


class XlsxExtractor(BaseExtractor):
    """Extrai valores de células de arquivos XLSX (modo somente leitura)."""

    extensions = ("xlsx", "xlsm")

    def is_available(self) -> bool:
        return _XLSX_OK

    def extract(self, path: Path, config: AppConfig) -> ExtractionResult:
        if not _XLSX_OK:
            return ExtractionResult(
                status="metadata_only",
                error="openpyxl não instalado (pip install openpyxl)",
            )
        workbook = None
        try:
            workbook = openpyxl.load_workbook(
                str(path), read_only=True, data_only=True
            )
            parts: list[str] = []
            total_cells = 0
            for sheet in workbook.worksheets:
                parts.append(str(sheet.title))
                for row_index, row in enumerate(sheet.iter_rows(values_only=True)):
                    if row_index >= _MAX_XLSX_ROWS_PER_SHEET:
                        break
                    values = [str(v) for v in row if v is not None]
                    total_cells += len(values)
                    if values:
                        parts.append(" | ".join(values))
                    if total_cells > _MAX_XLSX_CELLS:
                        break
                if total_cells > _MAX_XLSX_CELLS:
                    break
            return ExtractionResult(text="\n".join(parts)).truncated()
        except Exception as exc:
            logger.warning("Falha ao ler XLSX %s: %s", path, exc)
            return ExtractionResult(status="error", error=f"XLSX ilegível: {exc}")
        finally:
            if workbook is not None:
                workbook.close()
